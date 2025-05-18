import pandas as pd                      
import os                                
from docx import Document               
from docx2pdf import convert             
from zipfile import ZipFile 
import streamlit as st
import pythoncom       

def gerar_contratos(modelo_path, csv_path, output_dir):
    df = pd.read_csv(csv_path)

    os.makedirs(output_dir, exist_ok=True)

    for nome in df["nome"]:
        doc = Document(modelo_path) 

        for p in doc.paragraphs:
            if "NOME" in p.text:
                p.text = p.text.replace("NOME", nome)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "NOME" in cell.text:
                        cell.text = cell.text.replace("NOME", nome)

        caminho_docx = os.path.join(output_dir, f"Contrato - {nome}.docx")
        caminho_pdf = os.path.join(output_dir, f"Contrato - {nome}.pdf")

        doc.save(caminho_docx)
        pythoncom.CoInitialize() 
        convert(caminho_docx, caminho_pdf) 
        os.remove(caminho_docx)

    zip_path = os.path.join(output_dir, "contratos.zip")
    with ZipFile(zip_path, "w") as zipf:
        for nome in df["nome"]:
            caminho_pdf = os.path.join(output_dir, f"Contrato - {nome}.pdf")
            zipf.write(caminho_pdf, arcname=os.path.basename(caminho_pdf))  

    return zip_path

# Interface gráfica com Streamlit
st.title("Gerador de Contratos")


modelo_path = st.file_uploader("Selecione o modelo de contrato (.docx)", type=["docx"], key="modelo_uploader")
csv_path = st.file_uploader("Selecione o arquivo CSV com os dados", type=["csv"], key="csv_uploader")
output_dir = st.text_input("Diretório de saída", "contratos", key="output_dir_input")


if st.button("Gerar Contratos", key="gerar_contratos_button"):
    if modelo_path and csv_path and output_dir:
        try:
            modelo_temp = "modelo_temp.docx"
            csv_temp = "dados_temp.csv"
            with open(modelo_temp, "wb") as f:
                f.write(modelo_path.read())
            with open(csv_temp, "wb") as f:
                f.write(csv_path.read())

            zip_path = gerar_contratos(modelo_temp, csv_temp, output_dir)
            st.success(f"Contratos gerados com sucesso! Arquivo salvo em: {zip_path}")

            with open(zip_path, "rb") as f:
                st.download_button("Baixar arquivo ZIP", f, file_name="contratos.zip", key="download_zip_button")

            os.remove(modelo_temp)
            os.remove(csv_temp)
        except Exception as e:
            st.error(f"Erro ao gerar contratos: {e}")
    else:
        st.warning("Por favor, preencha todos os campos!")